import io
import json
import re
from flask import Flask, request, render_template_string, send_file, session, redirect, url_for
import PyPDF2
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)
app.secret_key = "your_secret_key_here"  # Replace with a secure key

VERSION = "1.19 (Web Version)"

# ----------------------------
# Allowed Size Orders per Category
# ----------------------------
size_order_map = {
    "Short Sleeve": ["XS", "S", "M", "L", "XL", "2XL", "3XL", "4XL"],
    "Short Sleeve V-Neck": ["XS", "S", "M", "L", "XL", "2XL", "3XL", "4XL"],
    "Tank Top": ["S", "M", "L", "XL", "2XL"],
    "Sweatshirt": ["S", "M", "L", "XL", "2XL", "3XL"],
    "Long Sleeve": ["S", "M", "L", "XL", "2XL", "3XL"],
    "Hoodie": ["S", "M", "L", "XL", "2XL", "3XL"],
    "Youth": ["S", "M", "L"],
    "Toddler": ["2T", "3T", "4T", "5T"],
    "Baby": ["NB", "0-6M", "6-12M", "12-18M", "18-24M"]
}

def get_size_order(size, category):
    allowed = size_order_map.get(category, [])
    try:
        return allowed.index(size.upper())
    except ValueError:
        return 100

# ----------------------------
# Color Normalization
# ----------------------------
def normalize_color(color_str):
    mapping = {
        "heather dark gray": "Heather Dark Grey",
        "heather dark grey": "Heather Dark Grey",
        "light pink": "Pink",
        "pink": "Pink",
        "sand": "Sand",
        "sand/natural": "Sand",
        "sand / natural": "Sand"
    }
    normalized = color_str.strip().lower()
    return mapping.get(normalized, color_str.strip().title())

# ----------------------------
# Extraction and Normalization Functions
# ----------------------------
def extract_orders_from_pdf(file_stream):
    try:
        reader = PyPDF2.PdfReader(file_stream)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        print("[INFO] Successfully extracted text from a PDF file.")
        return text
    except Exception as e:
        print(f"[ERROR] Failed to extract PDF text: {e}")
        return ""

def extract_items(text):
    pattern = r"(?:Quantity:\s*(\d+)).*?(?:(?:Select Shirt Size:? )|(?:Shirt Size:? )|(?:Size:? ))\s*([^\n]+).*?(?:(?:Select Shirt Color:? )|(?:Shirt Color:? )|(?:Color:? ))\s*([^\n]+)"
    matches = re.findall(pattern, text, re.DOTALL | re.IGNORECASE)
    items = []
    for qty, size, color in matches:
        try:
            qty = int(qty.strip())
        except:
            qty = 0
        items.append((qty, size.strip(), color.strip()))
    print(f"[INFO] Extracted {len(items)} items from text.")
    return items

def normalize_apparel_size(size_str):
    original = size_str.strip()
    ls = original.lower()
    # Baby category
    if "baby" in ls or "onesie" in ls:
        match = re.search(r'[-–—]\s*((?:NB)|(?:\d+\s*[-]\s*\d+M))', original, re.IGNORECASE)
        if match:
            size_value = match.group(1).replace(" ", "").upper()
            return f"Baby - {size_value}"
        if "nb" in ls:
            return "Baby - NB"
        match2 = re.search(r'(\d+\s*[-]\s*\d+M)', original, re.IGNORECASE)
        if match2:
            size_value = match2.group(1).replace(" ", "").upper()
            return f"Baby - {size_value}"
        return "Baby"
    # Youth category.
    if "youth" in ls:
        match = re.search(r'-\s*([\w]+)', original, re.IGNORECASE)
        if match:
            size_val = match.group(1).lower()
            if size_val in ["medium", "med"]:
                size_val = "M"
            else:
                size_val = size_val.upper()
            return f"Youth - {size_val}"
        return "Youth"
    # Toddler category.
    if "toddler" in ls:
        match = re.search(r'[-:]\s*([\d]+T)', original, re.IGNORECASE)
        if match:
            return f"Toddler - {match.group(1).upper()}"
        match2 = re.search(r'(\d+T)\b', original, re.IGNORECASE)
        if match2:
            return f"Toddler - {match2.group(1).upper()}"
        return "Toddler"
    # Adult apparel normalization:
    if "short sleeve" in ls:
        if "v-neck" in ls or "v neck" in ls or "v. neck" in ls or "vneck" in ls:
            match = re.search(r'(?:short sleeve.*?v[- ]?neck).*?[-–]\s*([\w\+]+)', original, re.IGNORECASE)
            if match:
                return f"Short Sleeve V-Neck - {match.group(1).upper()}"
            return "Short Sleeve V-Neck"
        else:
            match = re.search(r'short sleeve.*?[-–]\s*([\w\+]+)', original, re.IGNORECASE)
            if match:
                return f"Short Sleeve - {match.group(1).upper()}"
            return "Short Sleeve"
    if "tank" in ls:
        match = re.search(r'(?:tank(?: top)?).*?[-–]\s*([\w\+]+)', original, re.IGNORECASE)
        if match:
            return f"Tank Top - {match.group(1).upper()}"
        return "Tank Top"
    if "sweatshirt" in ls:
        match = re.search(r'sweatshirt.*?[-–]\s*([\w\+]+)', original, re.IGNORECASE)
        if match:
            return f"Sweatshirt - {match.group(1).upper()}"
        return "Sweatshirt"
    if "long sleeve" in ls:
        match = re.search(r'long sleeve.*?[-–]\s*([\w\+]+)', original, re.IGNORECASE)
        if match:
            return f"Long Sleeve - {match.group(1).upper()}"
        return "Long Sleeve"
    if "hoodie" in ls:
        match = re.search(r'hoodie.*?[-–]\s*([\w\+]+)', original, re.IGNORECASE)
        if match:
            return f"Hoodie - {match.group(1).upper()}"
        return "Hoodie"
    return original

def get_category_and_size(normalized):
    parts = normalized.split(" - ")
    if len(parts) >= 2:
        cat = parts[0].strip()
        size = parts[1].strip()
        if "v-neck" in cat.lower() or "v neck" in cat.lower() or "v. neck" in cat.lower():
            cat = "Short Sleeve V-Neck"
        return (cat, size)
    return (normalized, "")

def process_pdf_file(file_stream):
    text = extract_orders_from_pdf(file_stream)
    if not text:
        return None
    items = extract_items(text)
    if not items:
        return None
    df = pd.DataFrame(items, columns=["Quantity", "Shirt Size", "Shirt Color"])
    df["Shirt Color"] = df["Shirt Color"].apply(normalize_color)
    df["Normalized Size"] = df["Shirt Size"].apply(normalize_apparel_size)
    agg_df = df.groupby(["Normalized Size", "Shirt Color"], as_index=False)["Quantity"].sum()
    return agg_df

# ----------------------------
# DOCX Export Functionality (to BytesIO)
# ----------------------------
def set_repeat_table_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)

def export_to_docx_bytes(df):
    df[['Category', 'Size']] = df['Normalized Size'].apply(lambda x: pd.Series(get_category_and_size(x)))
    document = Document()
    for section in document.sections:
        section.top_margin = Inches(0.2)
        section.bottom_margin = Inches(0.2)
        section.left_margin = Inches(0.2)
        section.right_margin = Inches(0.2)
    style = document.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    available_width = (8.5 * 72) - (0.4 * 72)
    categories = sorted(df["Category"].unique())
    for cat in categories:
        cat_df = df[df["Category"] == cat]
        pivot = cat_df.pivot_table(index="Shirt Color", columns="Size", values="Quantity", fill_value=0, aggfunc='sum')
        pivot = pivot.reset_index()
        allowed = size_order_map.get(cat, sorted(list(pivot.columns)))
        size_columns = [col for col in pivot.columns if col != "Shirt Color"]
        size_columns.sort(key=lambda s: allowed.index(s.upper()) if s.upper() in allowed else 100)
        header_row = ["Color"] + size_columns
        num_cols = len(header_row)
        col_width = available_width / num_cols
        num_rows = 2 + len(pivot)
        table = document.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        first_hdr = table.rows[0].cells
        first_hdr[0].merge(first_hdr[-1])
        first_hdr[0].text = f"Category: {cat}"
        for paragraph in first_hdr[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
        set_repeat_table_header(table.rows[0])
        second_hdr = table.rows[1].cells
        for j, header in enumerate(header_row):
            second_hdr[j].text = header
            for paragraph in second_hdr[j].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
        set_repeat_table_header(table.rows[1])
        for i, row in pivot.iterrows():
            row_cells = table.rows[i+2].cells
            row_cells[0].text = str(row["Shirt Color"])
            for j, size in enumerate(size_columns):
                row_cells[j+1].text = str(row[size])
                for paragraph in row_cells[j+1].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(col_width / 72)
        document.add_paragraph()
    bytes_io = io.BytesIO()
    document.save(bytes_io)
    bytes_io.seek(0)
    return bytes_io

# ----------------------------
# Flask Routes
# ----------------------------
UPLOAD_FORM = """
<!doctype html>
<html>
<head>
  <title>Order Aggregation Report v{{version}}</title>
</head>
<body>
  <h1>Order Aggregation Report v{{version}}</h1>
  <form method="post" enctype="multipart/form-data">
    <label>Select PDF File(s):</label><br>
    <input type="file" name="pdf_files" multiple><br><br>
    <input type="submit" value="Upload and Process">
  </form>
  {% if table_html %}
    <h2>Aggregated Report</h2>
    {{ table_html|safe }}
    <form action="{{ url_for('download_csv') }}" method="post">
      <input type="submit" value="Download CSV">
    </form>
    <form action="{{ url_for('download_docx') }}" method="post">
      <input type="submit" value="Download DOCX">
    </form>
  {% endif %}
</body>
</html>
"""

@app.route("/", methods=["GET", "POST"])
def upload():
    table_html = ""
    if request.method == "POST":
        files = request.files.getlist("pdf_files")
        dfs = []
        for file in files:
            if file:
                file_stream = io.BytesIO(file.read())
                df = process_pdf_file(file_stream)
                if df is not None:
                    dfs.append(df)
        if dfs:
            combined_df = pd.concat(dfs, ignore_index=True)
            final_df = combined_df.groupby(["Normalized Size", "Shirt Color"], as_index=False)["Quantity"].sum()
            session["agg_data"] = final_df.to_json(orient="split")
            table_html = final_df.to_html(classes="table table-bordered", index=False)
        else:
            table_html = "<p>No order items found.</p>"
    return render_template_string(UPLOAD_FORM, version=VERSION, table_html=table_html)

@app.route("/download_csv", methods=["POST"])
def download_csv():
    agg_data = session.get("agg_data", None)
    if agg_data is None:
        return redirect(url_for("upload"))
    df = pd.read_json(agg_data, orient="split")
    csv_io = io.StringIO()
    df.to_csv(csv_io, index=False)
    csv_io.seek(0)
    return send_file(io.BytesIO(csv_io.getvalue().encode('utf-8')),
                     mimetype="text/csv",
                     as_attachment=True,
                     download_name="order_report.csv")

@app.route("/download_docx", methods=["POST"])
def download_docx():
    agg_data = session.get("agg_data", None)
    if agg_data is None:
        return redirect(url_for("upload"))
    df = pd.read_json(agg_data, orient="split")
    docx_io = export_to_docx_bytes(df)
    return send_file(docx_io,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     as_attachment=True,
                     download_name="order_report.docx")

if __name__ == "__main__":
    import os
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)